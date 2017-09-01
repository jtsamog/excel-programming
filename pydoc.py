from docx import Document

# define filename
filename = 'NE Competency 8 SC Certificate Assessment - new.docx'
filename2 = filename[0:filename.find('.docx')]+'3'+ filename[filename.find('.docx'):]

# read doc file
doc= Document(filename)
# to save
doc2 = Document()

Qs = []
q = []
for p in doc.paragraphs:
    line = p.text
    line = line.replace('\t',' ')
    highlight = False
    for r in p.runs:
        #print(r.font.highlight_color)
        if r.font.highlight_color:
            highlight = True
            break
    if highlight:
        line = '*' + line  
        
    # if it is empty line skip    
    if line == '':
        continue

    # if it is start with "Test Map:" skip
    if line.find("Test Map:") != -1:
        continue
    
    # check it the first character is number
    if line[0:1].isdigit():
        if q:
            Qs.append(q)
            q = []

    # if it is start with "Rationale:" do followings, then skip.
    #*** get the text after "Rationale:"
    #*** newLine1 = "~" + text
    #*** newLine2 = "@" + text
    #*** Insert these 2 lines after question
    if line.find("Rationale:") != -1:
        newline1 = line.replace('Rationale:','~')
        newline2 = line.replace('Rationale:','@')
        q.insert(1,newline2)
        q.insert(1,newline1)
        continue
    
    
    q.append(line)

    #print(line)
    #print(p.style)
    
Qs.append(q)

#print(q)
#print(Qs)

for q in Qs:
    for txt in q:
        doc2.add_paragraph(txt)

    # add a empty line after each question
    doc2.add_paragraph('')

doc2.save(filename2)
